import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('TripFlag Enterprise System Design, Architecture, & API Reference', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Confidential & Proprietary - Architecture Division\n', style='Subtitle')
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "TripFlag is a highly secure, enterprise-grade data ingestion and auditing platform designed for large-scale logistics and trip telemetry. "
        "It normalizes arbitrary file formats (XLSX, CSV, PDF, JPG) using advanced AI-driven Optical Character Recognition (OCR) and applies a rigorous "
        "deterministic rule-based flagging engine. The application architecture strictly separates concerns between the Edge-deployed Next.js web frontend, "
        "the native Electron desktop wrapper, and the high-performance Neon Serverless PostgreSQL database."
    )
    
    # 2. System Architecture
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        "The system employs a bifurcated client model. The Web Client (Next.js) is served globally via Vercel's "
        "Edge Network, providing standard portal access. The Desktop Client uses Electron to wrap the web application "
        "for local execution, enabling deeper OS integration while maintaining a unified codebase. Both clients communicate "
        "with the Node.js API Server via secure HTTPS/REST."
    )
    
    arch_img_path = 'docs/architecture_ieee.jpg'
    if os.path.exists(arch_img_path):
        doc.add_picture(arch_img_path, width=Inches(6.0))
        img_cap = doc.add_paragraph('Figure 1: High-Level System Architecture and Data Flow Boundaries.')
        img_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # 3. Database Schema
    doc.add_heading('3. Database Schema & Data Models', level=1)
    doc.add_paragraph(
        "The PostgreSQL database (hosted on Neon Serverless) relies on tightly normalized tables. The database leverages native JSONB "
        "columns to handle dynamic, unstructured row data extracted from unpredictable uploaded documents, while maintaining strict foreign keys."
    )
    
    db_img_path = 'docs/database_schema.jpg'
    if os.path.exists(db_img_path):
        doc.add_picture(db_img_path, width=Inches(6.0))
        img_cap = doc.add_paragraph('Figure 2: Entity Relationship Diagram (ERD) for Trips, Rows, and Rules.')
        img_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_heading('3.1 trips Table', level=2)
    doc.add_paragraph("Stores metadata for every file uploaded.\n"
                      "- id (UUID): Primary Key\n"
                      "- name, original_filename (VARCHAR): User-provided and original file names.\n"
                      "- file_type (VARCHAR): Extension mapping (xlsx, pdf, img).\n"
                      "- column_headers (JSONB): Array of extracted spreadsheet headers.\n"
                      "- total_rows, flagged_rows (INTEGER): Cached counts for O(1) reads.\n"
                      "- status (VARCHAR): 'pending', 'approved', 'rejected'.\n"
                      "- uploaded_at, approved_at (TIMESTAMP): Audit timestamps.")

    doc.add_heading('3.2 trip_rows Table', level=2)
    doc.add_paragraph("Stores individual row data. Normalizes spreadsheet grids and OCR output into JSONB objects.\n"
                      "- id (UUID): Primary Key\n"
                      "- trip_id (UUID): Foreign Key (trips.id), cascades on delete.\n"
                      "- row_index (INTEGER): Preserves original file row order.\n"
                      "- row_data (JSONB): Key-value pairing of the row data.\n"
                      "- is_flagged (BOOLEAN): High-speed filtering boolean.\n"
                      "- flag_details (JSONB): Array of specific rule violations.")
                      
    doc.add_heading('3.3 flagging_rules Table', level=2)
    doc.add_paragraph("Maintains the dynamic business logic parameters.\n"
                      "- id (UUID): Primary Key\n"
                      "- field_name (VARCHAR): Exact JSON key to evaluate in row_data.\n"
                      "- operator (VARCHAR): Logic constraint (greater_than, equals, contains, etc.).\n"
                      "- value, value_end (VARCHAR): Threshold values.\n"
                      "- unit (VARCHAR): Human readable metric.\n"
                      "- severity (VARCHAR): 'warning' or 'critical'.\n"
                      "- is_active (BOOLEAN): Soft-delete/toggle state.")

    # 4. API Reference
    doc.add_heading('4. Comprehensive API Reference', level=1)
    doc.add_paragraph(
        "The Node.js API layer is responsible for data sanitation, database transactions, and coordinating with AI microservices. "
        "All routes are implemented as Next.js API Routes (Serverless Functions)."
    )
    
    seq_img_path = 'docs/api_sequence_diagram.jpg'
    if os.path.exists(seq_img_path):
        doc.add_picture(seq_img_path, width=Inches(6.0))
        img_cap = doc.add_paragraph('Figure 3: API Sequence Diagram - Interaction Flow.')
        img_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('4.1 Trips API (/api/trips)', level=2)
    doc.add_paragraph(
        "GET /api/trips\n"
        "Description: Fetches all trip metadata, ordered by uploaded_at DESC.\n"
        "Response: JSON Array of Trip Objects (excluding raw row_data for bandwidth efficiency).\n\n"
        "POST /api/trips\n"
        "Description: Handles core data ingestion payload.\n"
        "Payload: { name: string, original_filename: string, file_type: string, column_headers: array, rows: array }\n"
        "Logic: \n"
        "1. Validates payload existence.\n"
        "2. Sanitizes JSONB to strip PostgreSQL-incompatible null-bytes (\\u0000).\n"
        "3. Inserts master 'trips' record.\n"
        "4. Iteratively inserts 'trip_rows' objects."
    )
    
    doc.add_heading('4.2 Trips Details API (/api/trips/[id])', level=2)
    doc.add_paragraph(
        "GET /api/trips/[id]\n"
        "Description: Fetches a deeply nested trip object, including all associated rows.\n"
        "Response: { id, name, ..., rows: [ { row_id, row_data, flag_details... } ] }\n\n"
        "PATCH /api/trips/[id]\n"
        "Description: Updates trip status.\n"
        "Payload: { status: 'approved' | 'rejected', name?: string }\n"
        "Logic: Updates status and writes to approved_at if status == 'approved'.\n\n"
        "DELETE /api/trips/[id]\n"
        "Description: Permanently deletes a trip. (Cascades to trip_rows)."
    )
    
    doc.add_heading('4.3 Rules API (/api/rules)', level=2)
    doc.add_paragraph(
        "GET /api/rules\n"
        "Description: Fetches all rules ordered by created_at DESC.\n\n"
        "POST /api/rules\n"
        "Description: Creates a new flagging rule.\n"
        "Payload: { field_name, operator, value, value_end, unit, severity, label, is_active }\n\n"
        "PATCH /api/rules\n"
        "Description: Toggles is_active status of a rule.\n"
        "Payload: { id, is_active: boolean }\n\n"
        "PUT /api/rules\n"
        "Description: Fully replaces rule thresholds and conditions.\n"
        "Payload: { id, field_name, operator, value... }\n\n"
        "DELETE /api/rules?id=[id]\n"
        "Description: Deletes a specific rule."
    )

    doc.add_heading('4.4 Flagging Execution API (/api/flag)', level=2)
    doc.add_paragraph(
        "POST /api/flag\n"
        "Description: Executes the deterministic flagging engine on a specific trip.\n"
        "Payload: { trip_id: string }\n"
        "Execution Flow:\n"
        "1. Queries 'trips' to verify existence.\n"
        "2. Queries all 'trip_rows' for the given trip_id.\n"
        "3. Queries all 'flagging_rules' where is_active = true.\n"
        "4. Executes 'runFlagging()' from lib/flagger.js locally in memory.\n"
        "5. Iteratively UPDATEs trip_rows setting is_flagged = true and injecting flag_details.\n"
        "6. Calculates total violations and UPDATEs trips.flagged_rows counter."
    )

    doc.add_heading('4.5 AI Integration APIs (/api/analyze-file-structure, /api/parse-pdf, /api/ocr)', level=2)
    doc.add_paragraph(
        "These endpoints act as proxies to the Groq Llama-3 AI microservices.\n"
        "- /api/analyze-file-structure: Sends the first 50 rows of a complex spreadsheet as a markdown table to the LLM to dynamically determine the integer index of the 'Header Row', avoiding strict templating constraints.\n"
        "- /api/parse-pdf & /api/ocr: Receives base64 image strings or parsed PDF text, prompting the LLM to act as a Tabular Data Extraction Engine. Bypasses traditional zonal OCR by relying on LLM semantic understanding to return strict JSON arrays."
    )
    
    # Save document
    doc.save('docs/TripFlag_Detailed_Architecture_And_API.docx')
    print("Document successfully generated at docs/TripFlag_Detailed_Architecture_And_API.docx")

if __name__ == '__main__':
    create_doc()
